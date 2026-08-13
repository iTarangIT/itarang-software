from docx import Document
from pathlib import Path

output_path = Path(r"C:\Users\Aniket\Downloads\iTarang-Notification-Flow.docx")

doc = Document()
doc.add_heading('iTarang Notification Flow - End to End', level=1)

doc.add_paragraph('This document describes the complete notification flow in the iTarang codebase, from event emission to bell delivery and notification access gating.')

doc.add_heading('1. Notification Type Catalog', level=2)
doc.add_paragraph('• Source: src/lib/notifications/catalog.ts')
doc.add_paragraph('• Maps every notification type to a category and determines whether the notification is email-worthy by default using emailWorthy(type).')
doc.add_paragraph('• Notification types are defined in CATEGORY_BY_TYPE and include both legacy flat types (e.g. loan_sanctioned, inventory_assigned) and modern dotted types (e.g. onboarding.agreement_signed).')
doc.add_paragraph('• The admin Notification Access screen derives its type list and labels from registry.ts, which imports this catalog and buysback notification metadata. New types appear automatically on the screen once added to the catalog.')

doc.add_heading('2. Notification Access / Bell Gating (E-231)', level=2)
doc.add_paragraph('• Source: src/lib/notifications/access.ts and src/app/api/admin/notification-access/route.ts')
doc.add_paragraph('• The feature controls in-app bell delivery only; email is unaffected by these settings.')
doc.add_paragraph('• notification_access rows with enabled = false are treated as muted. Absent rows mean enabled. The gate fails open on errors so no notification is silently dropped.')

doc.add_paragraph('• blockedDashboardsFor(type) loads denied dashboard/type pairs and returns the dashboards that have muted the type. It caches the result for 60 seconds and always returns an empty set on load failure.')
doc.add_paragraph('• filterAllowedRoles(roles, type) is used by notifyRoles to narrow a role-based audience before writing rows. blockedRoleClause(roleColumn, blocked) produces SQL filter clauses for INSERT … SELECT paths used by notifyNbfcTenant and notifyRoles.')

doc.add_paragraph('• The admin API route GET returns only dashboards editable by the current user and the DENIAL rows. The PATCH route validates dashboard scope, known notification types, upserts notification_access rows, writes an audit log, and calls invalidateNotificationAccessCache().')

doc.add_heading('3. Core Emit Function', level=2)
doc.add_paragraph('• Source: src/lib/notifications/emit.ts')
doc.add_paragraph('• emit(input) is the CRM-wide notification emitter. It is the central place where one business event is fanned out to multiple recipients and optionally emailed.')

doc.add_paragraph('• emit() accepts a type, title, message, optional leadId/stage/from, a recipient list, shared data, and email overrides.')

doc.add_paragraph('• Each recipient has an audience, address party, optional href, actions, per-recipient title/message, email override, and extra data.')

doc.add_paragraph('• Recipients are resolved to concrete user logins via resolve(audience). The supported audience kinds are:')
doc.add_paragraph('  - roles: all active users with one of the specified roles')
doc.add_paragraph('  - user: a specific userId')
doc.add_paragraph('  - dealer: every active login for a dealer code')
doc.add_paragraph('  - lead_dealer: the dealer of a lead, resolved through leads.dealer_id')
doc.add_paragraph('  - nbfc: every seat of one NBFC tenant')
doc.add_paragraph('  - lead_nbfc: every NBFC assigned to a lead, optionally excluding one tenant')

doc.add_paragraph('• Dedupe ensures one notification per user_id even if the same person is reachable from multiple audiences.')

doc.add_paragraph('• The bell gate is applied after deduplication and before writing rows. blockedDashboardsFor(type) is used to drop muted dashboard roles from bellTargets, while the email path still runs for the original targets.')

doc.add_paragraph('• If any bellTargets remain, emit inserts one notifications row per target with id, user_id, dealer_id (for backward compatibility), lead_id, type, title, message, data, and read=false.')

doc.add_paragraph('• Email is sent when recipient.email is true or the type is email-worthy by default. emailTargets() collects valid email addresses and sends either through the platform transport or an NBFC tenant-specific transport when tenantId is present.')

doc.add_paragraph('• emit() never throws. Errors in one audience are logged and do not stop other audiences or the caller.')

doc.add_heading('4. Helper Notification Writers', level=2)
doc.add_paragraph('• Source: src/lib/notifications/notify.ts')
doc.add_paragraph('• notifyRoles(roles, n) is a lightweight helper that writes one notification row to every active user holding the requested roles.')
doc.add_paragraph('• notifyRoles uses filterAllowedRoles() before the INSERT to apply the E-231 gate to role-based audiences. It then inserts directly into notifications with JSONB data payload.')

doc.add_paragraph('• notifyUser(userId, n) resolves one user by ID, checks the user role against blockedDashboardsFor(type), and inserts one row if not muted. This preserves current behaviour for missing user rows by not folding the guard into the INSERT.')

doc.add_paragraph('• notifyNbfcTenant(tenantId, n) inserts one row for each active nbfc_users seat of the tenant, joining through users so the bell gate uses the CRM dashboard role and not the finer NBFC seat role.')

doc.add_heading('5. Notification Registry and Admin Screen', level=2)
doc.add_paragraph('• Source: src/lib/notifications/registry.ts')
doc.add_paragraph('• The registry defines the dashboard metadata, editable dashboard scope, and human labels for every notification type. It also imports buyback notification metadata for buyback events.')

doc.add_paragraph('• editableDashboardsFor(role) defines who may change notification access: admin/ceo can edit all dashboards; sales_head can edit its sales tree; others cannot edit.')

doc.add_paragraph('• The admin Notification Access page renders checkboxes for the dashboard/type matrix using this registry and the catalog-derived type list. Unknown roles are surfaced in the API response so the screen does not silently claim full coverage.')

doc.add_heading('6. Practical Side Effects', level=2)
doc.add_paragraph('• A muted notification type is silent only in the in-app bell. The email path still sends unless email is explicitly disabled by the emitter or recipient override.')
doc.add_paragraph('• The system fails open on notification access load failure, so missing or broken E-231 state does not stop business events from producing notifications.')

doc.add_paragraph('• New notification types are default-on unless they are also added to the registry labels and then explicitly muted by an admin. Unmapped types remain always-on and are not governed by the access screen.')

doc.add_paragraph('• The bell requires user_id on notifications rows; older rows with only dealer_id are invisible to the per-user feed. emit() writes user_id and dealer_id to maintain compatibility.')

output_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(output_path)
print(str(output_path))
