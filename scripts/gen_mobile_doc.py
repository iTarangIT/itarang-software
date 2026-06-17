# -*- coding: utf-8 -*-
"""SIMPLE 'Login + Dealer Dashboard Mobile-Friendly' document (.docx).
System design is shown as SEPARATE standalone boxes (no flow chart)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x02, 0x31, 0x4E)
SKY = RGBColor(0x13, 0x8F, 0xC6)
GREY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x1F, 0x29, 0x37)
CODE_BG = "F2F4F6"

doc = Document()
for s in doc.sections:
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = DARK


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def plan_box(title, desc, fill, text_rgb, tag=None):
    """One standalone colored box (a card) with a title and a short line."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(tbl)
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    if tag:
        rt = p.add_run(tag + "   ")
        rt.bold = True
        rt.font.size = Pt(9)
        rt.font.color.rgb = text_rgb
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = text_rgb
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(5)
    r2 = p2.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = text_rgb
    doc.add_paragraph()  # gap = boxes look separate


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level <= 2 else SKY
    return p


def para(text, size=11):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size); return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True; p.add_run(it[1])
        else:
            p.add_run(it)


def code_block(text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0); shade(cell, CODE_BG); cell.text = ""
    for i, line in enumerate(text.strip("\n").split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line); r.font.name = "Consolas"; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    doc.add_paragraph()


def simple_table(rows, headers, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        tbl.rows[0].cells[i].text = hd
        for pr in tbl.rows[0].cells[i].paragraphs:
            for run in pr.runs:
                run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    doc.add_paragraph()


# ===================================================== TITLE
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(10)
r = t.add_run("Making the Login Page & Dealer Dashboard Mobile-Friendly")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("A Simple Guide  ·  Mobile screen only  ·  Desktop UI, Database & Backend stay the same")
r.font.size = Pt(12); r.font.color.rgb = SKY
doc.add_paragraph()

# ===================================================== 1. THE IDEA
h("1. What We Want To Do", 1)
para("Today the app looks good on a computer but is hard to use on a phone. We want the "
     "Login page and the Dealer Dashboard to also work nicely on phones.")
para("We do this WITHOUT changing the desktop look. The same page simply rearranges itself "
     "to fit the screen size. We do NOT touch the database or the backend - only the screen "
     "(frontend).")

# ===================================================== 2. SYSTEM DESIGN (BOXES)
h("2. System Design - The Plan in Boxes", 1)
para("Here is the whole plan as separate boxes. YELLOW boxes are the parts we change "
     "(only the screen). GREEN boxes stay exactly the same. Each box stands on its own.")
doc.add_paragraph()

py = RGBColor(0x92, 0x5A, 0x00)   # text color for yellow boxes
pg = RGBColor(0x1B, 0x5E, 0x20)   # text color for green boxes

plan_box("Login Page",
         "Make the login form fit a phone and scroll if needed. The big side picture stays "
         "hidden on phones (already does). File: app/(auth)/login/page.tsx",
         "FFE9B3", py, tag="[ CHANGE - SMALL ]")

plan_box("Mobile Menu (NEW)",
         "Add a hamburger button + a slide-out menu so a dealer can move between sections on "
         "a phone. Files: header.tsx, sidebar.tsx, new MobileNavDrawer.tsx, new uiStore.ts",
         "FFE9B3", py, tag="[ CHANGE + NEW ]")

plan_box("Dashboard & List Pages",
         "Stat cards stack one per row. Wide tables (Leads, Inventory, Orders...) become "
         "easy-to-read cards on phones. Files: dealer-portal pages + DealerDashboard.tsx",
         "FFE9B3", py, tag="[ CHANGE ]")

plan_box("Pop-up Windows (Modals)",
         "Make pop-ups fit the screen and scroll when they are long. Files: dealer-portal modals",
         "FFE9B3", py, tag="[ CHANGE ]")

plan_box("Backend & APIs (Next.js routes)",
         "No change. Same endpoints (/api/dealer/..., /api/user/profile), same data in and out.",
         "D6F0D6", pg, tag="[ NO CHANGE ]")

plan_box("Database (PostgreSQL on AWS) + Queue",
         "No change. Same tables, same rows, same Redis/BullMQ background jobs.",
         "D6F0D6", pg, tag="[ NO CHANGE ]")

plan_box("Desktop UI",
         "No change. On a computer everything looks exactly the same as it does today.",
         "D6F0D6", pg, tag="[ NO CHANGE ]")

# ===================================================== 3. THE GOLDEN RULE
h("3. The Golden Rule: Don't Break the Desktop", 1)
para("We use Tailwind CSS, which is 'mobile-first'. This gives us a safe trick so the desktop "
     "never changes:")
bullets([
    ("Plain class = phone: ", "a class with no prefix (e.g. grid-cols-1) sets the phone layout."),
    ("md: / lg: = bigger screens: ", "a class like md:grid-cols-4 only works on tablets/computers "
     "(768px and up). The desktop already uses these, so it stays the same."),
    ("md:hidden = phone only: ", "new mobile pieces (slide-out menu, hamburger) get md:hidden, so "
     "they NEVER appear on desktop."),
    ("hidden md:block = desktop only: ", "the existing desktop sidebar already has this, so it "
     "stays exactly as it is."),
])
para("Meaning: we only ADD phone styles and a phone-only menu. We do not remove or rewrite the "
     "desktop styles.")
code_block(
    "EXAMPLE - one line of code serves both screens:\n\n"
    '   <div className="grid grid-cols-1 sm:grid-cols-2 lg:grid-cols-4 gap-4">\n\n'
    "   Phone  (grid-cols-1)     -> 1 box per row\n"
    "   Tablet (sm:grid-cols-2)  -> 2 boxes per row\n"
    "   Desktop(lg:grid-cols-4)  -> 4 boxes per row  (SAME AS NOW)"
)

# ===================================================== 4. LOGIN PAGE
h("4. Login Page - What To Change", 1)
para("The login page is already mostly responsive. It only needs small touches so it feels "
     "comfortable on a phone.")
simple_table(
    [
        ("src/app/(auth)/login/page.tsx",
         "Let the form scroll on short screens (overflow-y-auto + py-8); make the heading smaller "
         "on phones (text-2xl sm:text-3xl); keep the side image as hidden lg:flex so desktop is "
         "unchanged."),
    ],
    headers=("File", "What to change"),
    widths=[Inches(2.6), Inches(4.2)],
)
para("How it will look on a phone:")
code_block(
    "  +---------------------------+\n"
    "  |       [ iTarang logo ]    |\n"
    "  |   Welcome to iTarang      |\n"
    "  |   EVERY THING EV!         |\n"
    "  |   Email address           |\n"
    "  |   [____________________]  |\n"
    "  |   Password                |\n"
    "  |   [_______________] (eye) |\n"
    "  |   [     Sign In      ]    |\n"
    "  |   Don't have an account?  |\n"
    "  +---------------------------+\n"
    "  (The big rickshaw picture shows only on wide screens.)"
)

# ===================================================== 5. DEALER DASHBOARD
h("5. Dealer Dashboard - What To Change", 1)

h("5A. The Menu & Header (most important)", 2)
para("Problem: on a phone the left menu disappears and there is no button to open it, so the "
     "dealer gets stuck. Fix: add a phone-only slide-out menu and a hamburger button.")
simple_table(
    [
        ("src/components/layout/sidebar.tsx", "Change",
         "Move the menu list into a small reusable piece (SidebarNav) so the same links work in "
         "both the desktop sidebar and the phone menu. Keep 'hidden md:flex' so the DESKTOP "
         "sidebar is unchanged."),
        ("src/components/layout/MobileNavDrawer.tsx", "NEW",
         "A phone-only (md:hidden) menu that slides in from the left. Shows the same links. "
         "Closes on tap, on the dark background, or on page change."),
        ("src/components/layout/header.tsx", "Change",
         "Add a hamburger (three-line) button on the left, marked md:hidden, that opens the "
         "slide-out menu. Desktop header unchanged."),
        ("src/components/layout/LayoutWrapper.tsx", "Change",
         "Show the new drawer in the layout. (It already removes the left margin on phones.)"),
        ("src/store/uiStore.ts", "NEW (small)",
         "A tiny on/off switch (open / close) shared by the hamburger button and the drawer. "
         "Uses Zustand, which the project already has."),
    ],
    headers=("File", "Type", "What to do"),
    widths=[Inches(2.5), Inches(0.9), Inches(3.4)],
)

h("5B. The Pages (rearrange for small screens)", 2)
simple_table(
    [
        ("src/components/dealer-dashboard/DealerDashboard.tsx",
         "Stat cards: grid-cols-1 sm:grid-cols-2 lg:grid-cols-4 so they stack on phones. Loan "
         "cards stack too."),
        ("src/app/(dashboard)/dealer-portal/leads/page.tsx",
         "Fix side-by-side grid (grid-cols-2 -> grid-cols-1 sm:grid-cols-2). Turn the leads table "
         "into cards on phones (hidden md:table + a md:hidden card list)."),
        ("src/app/(dashboard)/dealer-portal/leads/drafts/page.tsx",
         "Same table-to-cards change for the drafts list."),
        ("src/app/(dashboard)/dealer-portal/inventory/page.tsx",
         "Table to cards; search box full width on phones."),
        ("src/app/(dashboard)/dealer-portal/orders/page.tsx", "Table to cards."),
        ("assets / batteries / service / loans pages",
         "Same pattern: tables become cards; chips and buttons wrap instead of overflowing."),
        ("All dealer pop-up windows (modals)",
         "Add max-h-[90vh] overflow-y-auto and p-4 sm:p-6 so they fit a phone."),
        ("Lead creation wizard",
         "Stack the steps and the two-column form rows on phones; Next/Back buttons full width."),
    ],
    headers=("File / Area", "What to change"),
    widths=[Inches(3.1), Inches(3.7)],
)

# ===================================================== 6. HOW IT LOOKS
h("6. How The Mobile Design Will Look", 1)
para("Dashboard with the menu CLOSED. A hamburger button appears at the top:")
code_block(
    "  +-----------------------------+\n"
    "  | [=]  iTarang        (you)   |   <- [=] = hamburger button (phones only)\n"
    "  +-----------------------------+\n"
    "  |  Total Leads        128     |   <- stat cards stack one per row\n"
    "  +-----------------------------+\n"
    "  |  Converted           34     |\n"
    "  +-----------------------------+\n"
    "  |  Recent Leads               |\n"
    "  |  +-----------------------+  |   <- table row shown as a card\n"
    "  |  | Ramesh K.  · New      |  |\n"
    "  |  | 98xxxxxx · Finance    |  |\n"
    "  |  +-----------------------+  |\n"
    "  +-----------------------------+"
)
para("Dashboard with the menu OPEN (after tapping the hamburger). It slides in from the left:")
code_block(
    "  +-----------------------------+\n"
    "  | iTarang            [x]      |\n"
    "  |-----------------------------|\n"
    "  |  Dashboard                  |\n"
    "  |  Lead Management            |   <- same links as the desktop sidebar\n"
    "  |  My Drafts                  |\n"
    "  |  Loan Processing            |\n"
    "  |  Inventory                  |\n"
    "  |  Orders from OEM            |\n"
    "  +-----------------------------+\n"
    "   (Tap a link, the [x], or the dark area to close.)"
)
para("On a computer: nothing changes. The full sidebar stays on the left and the hamburger "
     "button never appears.")

# ===================================================== 7. BEFORE / AFTER
h("7. Before and After (on a phone)", 1)
simple_table(
    [
        ("Login page", "Works, heading a bit large", "Comfortable, scrolls if needed"),
        ("Menu / navigation", "Gone - dealer is stuck", "Slide-out menu via a button"),
        ("Stat cards", "Squeezed side-by-side", "Stack neatly, one per row"),
        ("Wide tables", "Hard to read, cut off", "Shown as easy-to-read cards"),
        ("Pop-up windows", "Too tall, cut off", "Fit the screen and scroll"),
        ("Desktop look", "(good)", "(unchanged - exactly the same)"),
        ("Database & backend", "(unchanged)", "(unchanged - exactly the same)"),
    ],
    headers=("Item", "Now", "After the change"),
)

# ===================================================== 8. SUMMARY
h("8. Summary", 1)
para("We make the Login page and Dealer Dashboard work on phones by changing only the "
     "frontend. The main work is the phone menu (new MobileNavDrawer + hamburger) and "
     "rearranging the pages with Tailwind's md:/lg: rules. Because we only ADD phone styles "
     "and a phone-only menu, the desktop UI, the database, and all backend APIs stay exactly "
     "the same.")

out = r"C:\Users\Aniket\OneDrive\Desktop\Itarang_17_04_2026\itarang-software\docs\Dealer_Dashboard_Mobile_Friendly_Design.docx"
doc.save(out)
print("Saved:", out)
