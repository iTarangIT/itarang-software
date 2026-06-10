"""
Generate a simple-language Word (.docx) document describing the FUTURE
WhatsApp-based dealer onboarding chatbot for iTarang, with an embedded flowchart.

Run:  python scripts/docs/gen_whatsapp_onboarding_doc.py

Output: C:\\Users\\Aniket\\OneDrive\\Desktop\\WhatsApp_Dealer_Onboarding_Plan.docx

This is a one-off documentation generator. It does NOT touch the application code.
Requires: python-docx, matplotlib (both confirmed installed).
"""

import os
import tempfile

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Polygon

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
DESKTOP = r"C:\Users\Aniket\OneDrive\Desktop"
OUTPUT_DOCX = os.path.join(DESKTOP, "WhatsApp_Dealer_Onboarding_Plan.docx")

# Colors (role-coded, used in both the flowchart and the doc accents)
C_DEALER = "#25D366"   # WhatsApp green  -> dealer
C_BOT    = "#128C7E"   # teal            -> bot / system
C_EXTRACT= "#5B8DEF"   # blue            -> extraction service
C_ADMIN  = "#7E57C2"   # purple          -> admin
C_PAY    = "#F5A623"   # orange          -> payment
C_DONE   = "#2E7D32"   # dark green      -> approved
C_NO     = "#E53935"   # red             -> resend / rejected


# ---------------------------------------------------------------------------
# 1) Draw the flowchart with matplotlib and save as PNG
# ---------------------------------------------------------------------------
def build_flowchart(png_path: str) -> None:
    fig, ax = plt.subplots(figsize=(8.0, 11.0), dpi=160)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 26)
    ax.axis("off")

    def box(cx, cy, text, color, w=5.0, h=1.6, text_color="white", fontsize=11):
        x = cx - w / 2
        y = cy - h / 2
        patch = FancyBboxPatch(
            (x, y), w, h,
            boxstyle="round,pad=0.06,rounding_size=0.25",
            linewidth=1.2, edgecolor="#333333", facecolor=color,
        )
        ax.add_patch(patch)
        ax.text(cx, cy, text, ha="center", va="center",
                color=text_color, fontsize=fontsize, fontweight="bold", wrap=True)

    def diamond(cx, cy, text, color, w=6.2, h=2.6, fontsize=10):
        pts = [(cx, cy + h / 2), (cx + w / 2, cy), (cx, cy - h / 2), (cx - w / 2, cy)]
        ax.add_patch(Polygon(pts, closed=True, linewidth=1.2,
                             edgecolor="#333333", facecolor=color))
        ax.text(cx, cy, text, ha="center", va="center",
                color="white", fontsize=fontsize, fontweight="bold")

    def arrow(x1, y1, x2, y2, label=None, label_dx=0.0, label_dy=0.0):
        ax.add_patch(FancyArrowPatch(
            (x1, y1), (x2, y2),
            arrowstyle="-|>", mutation_scale=16,
            linewidth=1.6, color="#333333",
        ))
        if label:
            ax.text((x1 + x2) / 2 + label_dx, (y1 + y2) / 2 + label_dy, label,
                    ha="center", va="center", fontsize=9.5,
                    fontweight="bold", color="#333333",
                    bbox=dict(boxstyle="round,pad=0.18", fc="white", ec="#cccccc"))

    cx = 5.0
    # Vertical layout, top to bottom
    box(cx, 25.0, "Dealer says \"Hi\" on WhatsApp", C_DEALER)
    arrow(cx, 24.2, cx, 23.4)

    box(cx, 22.6, "Bot asks for documents\n(PAN, Aadhaar, GST, Bank)", C_BOT)
    arrow(cx, 21.8, cx, 21.0)

    box(cx, 20.2, "Dealer sends photos of\nthe documents", C_DEALER)
    arrow(cx, 19.4, cx, 18.6)

    box(cx, 17.8, "Extraction service reads\nthe data from photos", C_EXTRACT)
    arrow(cx, 17.0, cx, 16.2)

    box(cx, 15.4, "Bot shows the data back,\ndealer confirms it", C_BOT)
    arrow(cx, 14.6, cx, 13.8)

    box(cx, 13.0, "Data sent to Admin\nfor review", C_ADMIN)
    arrow(cx, 12.2, cx, 11.4)

    box(cx, 10.6, "Admin sends payment QR,\ndealer pays the fee", C_PAY)
    arrow(cx, 9.8, cx, 8.7)

    diamond(cx, 7.2, "Documents OK\nAND payment received?", C_ADMIN)

    # NO branch -> right then up back to "send photos"
    arrow(cx + 3.1, 7.2, 9.2, 7.2, label="NO", label_dy=0.5)
    ax.add_patch(FancyArrowPatch((9.2, 7.2), (9.2, 20.2),
                 arrowstyle="-|>", mutation_scale=16, linewidth=1.6,
                 color=C_NO, connectionstyle="arc3,rad=0.0"))
    arrow(9.2, 20.2, cx + 2.5, 20.2, label="Ask dealer to resend", label_dy=0.5)

    # YES branch -> down
    arrow(cx, 5.9, cx, 5.1, label="YES", label_dx=0.6)
    box(cx, 4.3, "Admin approves the dealer", C_DONE)
    arrow(cx, 3.5, cx, 2.7)
    box(cx, 1.9, "Dealer gets login +\nwelcome message", C_DEALER)

    # Legend
    legend_items = [
        ("Dealer", C_DEALER), ("Bot / System", C_BOT),
        ("Extraction", C_EXTRACT), ("Admin", C_ADMIN),
        ("Payment", C_PAY), ("Approved", C_DONE),
    ]
    lx = 0.2
    for i, (lab, col) in enumerate(legend_items):
        yy = 0.7 - 0.0  # baseline row
        ax.add_patch(FancyBboxPatch((lx + i * 1.62, 0.2), 0.3, 0.3,
                     boxstyle="round,pad=0.02", facecolor=col, edgecolor="#333"))
        ax.text(lx + i * 1.62 + 0.38, 0.35, lab, fontsize=7.5, va="center")

    ax.set_title("WhatsApp Dealer Onboarding — Flowchart",
                 fontsize=14, fontweight="bold", pad=12)

    fig.tight_layout()
    fig.savefig(png_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# ---------------------------------------------------------------------------
# 1b) Draw a simple timeline (Gantt-style) chart and save as PNG
# ---------------------------------------------------------------------------
def build_timeline(png_path: str) -> None:
    # (label, start_week, duration_weeks, color)
    phases = [
        ("Phase 1: Setup & WhatsApp approvals", 0, 2, C_DEALER),
        ("Phase 2: Bot + webhook + uploads",    1, 2, C_BOT),
        ("Phase 3: Extraction + verification",  3, 2, C_EXTRACT),
        ("Phase 4: Payment QR + admin approval", 4, 2, C_PAY),
        ("Phase 5: Testing + pilot",            6, 2, C_ADMIN),
    ]
    fig, ax = plt.subplots(figsize=(9.0, 4.2), dpi=160)
    for i, (label, start, dur, color) in enumerate(phases):
        y = len(phases) - i - 1
        ax.barh(y, dur, left=start, height=0.55, color=color,
                edgecolor="#333333", linewidth=1.0)
        ax.text(start + dur / 2, y, f"{dur} wk", ha="center", va="center",
                color="white", fontsize=10, fontweight="bold")
        ax.text(-0.2, y, label, ha="right", va="center", fontsize=10)

    ax.set_xlim(0, 8)
    ax.set_ylim(-0.6, len(phases) - 0.3)
    ax.set_yticks([])
    ax.set_xticks(range(0, 9))
    ax.set_xticklabels([f"W{w}" for w in range(0, 9)])
    ax.set_xlabel("Weeks from start", fontsize=10, fontweight="bold")
    ax.set_title("Project Timeline (about 8 weeks total)",
                 fontsize=13, fontweight="bold", pad=10)
    ax.grid(axis="x", linestyle="--", alpha=0.4)
    for spine in ["top", "right", "left"]:
        ax.spines[spine].set_visible(False)

    fig.tight_layout()
    fig.savefig(png_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# ---------------------------------------------------------------------------
# 2) Build the Word document
# ---------------------------------------------------------------------------
def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    return h


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p


def build_doc(png_path: str, timeline_png: str) -> None:
    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ----- Title page -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("WhatsApp Dealer Onboarding")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = RGBColor.from_string("128C7E")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("A Simple Plan (Documentation + Flowchart)")
    sr.font.size = Pt(14)
    sr.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("iTarang CRM  •  Date: 09 June 2026  •  Status: Future Plan / Idea").font.size = Pt(11)

    doc.add_paragraph()

    # ----- 1. What we are building -----
    add_heading(doc, "1. What we are building", level=1, color=C_BOT)
    add_para(doc, "We want a WhatsApp chatbot for dealer onboarding. Instead of opening a "
                  "website, a new dealer can simply chat on WhatsApp.")
    add_para(doc, "The dealer sends photos of their documents on WhatsApp. The bot reads the "
                  "data from those documents automatically. The data is sent to the admin. "
                  "The admin sends a payment QR code. The dealer pays. After the documents are "
                  "verified and the payment is received, the admin approves the dealer. The "
                  "dealer then gets a login and a welcome message.")
    add_para(doc, "In short: WhatsApp is the new, easy front door. Behind it, we reuse the "
                  "systems iTarang already has.")

    # ----- 2. The two services we need -----
    add_heading(doc, "2. The two services we need", level=1, color=C_BOT)
    add_para(doc, "To build this idea, we need two outside services.")

    add_heading(doc, "Service 1 — WhatsApp (to chat with the dealer)", level=2, color=C_DEALER)
    add_para(doc, "Recommended: Meta WhatsApp Cloud API (the official WhatsApp Business API "
                  "from Meta).", bold=True)
    add_para(doc, "What it does:")
    for b in [
        "Sends and receives WhatsApp messages with the dealer.",
        "Receives the document photos the dealer uploads.",
        "Sends the payment QR image and the final welcome message.",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    add_para(doc, "Why this choice: it is the official API, it is free to use (you only pay "
                  "small per-conversation fees to Meta), and it gives us full control to "
                  "connect it to our own Next.js backend.")
    add_para(doc, "What we need to set up:")
    for b in [
        "A Meta Business account and a verified business.",
        "A WhatsApp phone number for the bot.",
        "Message templates approved by Meta (for the first message we send).",
        "A webhook URL (an endpoint in our app that receives incoming messages).",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    add_heading(doc, "Service 2 — Document extraction (to read the documents)", level=2, color=C_EXTRACT)
    add_para(doc, "This service reads the data (name, number, address, IFSC, etc.) out of the "
                  "PAN, Aadhaar, GST, and bank photos the dealer sends. We compared two options:")

    # Comparison table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Point"
    hdr[1].text = "Option A: Reuse what we have (Tesseract + Decentro)"
    hdr[2].text = "Option B: Cloud AI (Google Document AI / AWS Textract)"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
    rows = [
        ("Cost", "Cheapest. No new vendor.", "Has a per-document cost."),
        ("Accuracy", "Medium. Struggles with blurry WhatsApp photos.", "High. Built for ID and financial documents."),
        ("Setup effort", "Already built in our app.", "New vendor account + integration."),
        ("Best for", "Clean, clear scans.", "Messy mobile photos (common on WhatsApp)."),
    ]
    for a, b, c in rows:
        cells = table.add_row().cells
        cells[0].text = a
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = b
        cells[2].text = c

    doc.add_paragraph()
    add_para(doc, "Recommendation:", bold=True)
    add_para(doc, "Use Cloud AI (Google Document AI) to READ the data, because WhatsApp photos "
                  "are often blurry and need higher accuracy. Then reuse our existing Decentro "
                  "KYC service to VERIFY the data (check PAN, Aadhaar, and bank account are "
                  "genuine). This gives us the best of both: accurate reading + trusted "
                  "verification.")

    # ----- 3. Step-by-step flow -----
    add_heading(doc, "3. Step-by-step flow (simple words)", level=1, color=C_BOT)
    steps = [
        "The dealer sends \"Hi\" to our WhatsApp number.",
        "The bot replies and asks for the documents (PAN, Aadhaar, GST, Bank).",
        "The dealer sends photos of each document.",
        "The extraction service reads the data from the photos.",
        "The bot shows the data back to the dealer and asks them to confirm it is correct.",
        "The confirmed data is sent to the admin for review.",
        "The admin reviews the documents and the data.",
        "The admin sends a payment QR code on WhatsApp.",
        "The dealer scans the QR and pays the fee.",
        "The system checks: are the documents OK and is the payment received?",
        "If something is wrong, the bot asks the dealer to send the document again.",
        "If everything is fine, the admin approves the dealer.",
        "The dealer receives a login and a welcome message on WhatsApp.",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")

    # ----- 4. The Flowchart -----
    doc.add_page_break()
    add_heading(doc, "4. The Flowchart", level=1, color=C_BOT)
    add_para(doc, "The picture below shows the full flow from top to bottom.")
    doc.add_picture(png_path, width=Inches(6.0))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "Text version of the flowchart (in case the picture does not show):", bold=True)
    text_flow = [
        "Dealer says Hi  ->  Bot asks for documents  ->  Dealer sends photos",
        "->  Extraction service reads data  ->  Bot shows data, dealer confirms",
        "->  Data sent to Admin  ->  Admin sends payment QR, dealer pays",
        "->  CHECK: Documents OK AND payment received?",
        "      - NO  ->  Ask dealer to resend the document (go back)",
        "      - YES ->  Admin approves dealer  ->  Dealer gets login + welcome message",
    ]
    for t in text_flow:
        p = doc.add_paragraph(t)
        p.runs[0].font.name = "Consolas"
        p.runs[0].font.size = Pt(10)

    # ----- 5. What connects to what we already have -----
    doc.add_page_break()
    add_heading(doc, "5. How it connects to what iTarang already has", level=1, color=C_BOT)
    add_para(doc, "The WhatsApp bot is only a new front door. The real work reuses systems we "
                  "already built. This keeps the project small and fast.")
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Light Grid Accent 1"
    h2 = t2.rows[0].cells
    h2[0].text = "Step in the flow"
    h2[1].text = "What it reuses"
    h2[2].text = "Where it lives in the code"
    for c in h2:
        c.paragraphs[0].runs[0].bold = True
    conn = [
        ("Read document data", "Our OCR (or new Cloud AI)", "src/lib/ocr/tesseractOcr.ts, src/lib/ocr/bankDocParser.ts"),
        ("Verify PAN / Aadhaar / Bank", "Decentro KYC", "src/lib/decentro.ts, src/lib/kyc/pan-verification.ts, src/lib/kyc/bank-verification.ts"),
        ("Store the application + docs", "Existing onboarding tables", "dealerOnboardingApplications, dealerOnboardingDocuments (src/lib/db/schema.ts)"),
        ("Payment QR for the fee", "Razorpay UPI QR", "src/lib/razorpay.ts -> createPaymentQr()"),
        ("Admin approves dealer", "Existing approve route", "src/app/api/admin/dealer-verifications/[dealerId]/approve/route.ts"),
    ]
    for a, b, c in conn:
        cells = t2.add_row().cells
        cells[0].text = a
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = b
        cells[2].text = c
        for run in cells[2].paragraphs[0].runs:
            run.font.size = Pt(9)

    # ----- 6. Checklist -----
    doc.add_paragraph()
    add_heading(doc, "6. What to set up next (checklist)", level=1, color=C_BOT)
    checklist = [
        "Create a Meta Business account and verify the business.",
        "Get a WhatsApp number and connect the WhatsApp Cloud API.",
        "Create and submit the first message templates for Meta approval.",
        "Build a webhook endpoint in our app to receive WhatsApp messages.",
        "Choose the extraction service (recommended: Google Document AI) and connect it.",
        "Reuse Decentro to verify the extracted PAN / Aadhaar / Bank data.",
        "Reuse the Razorpay QR (createPaymentQr) to collect the dealer fee.",
        "Send the data to the admin and reuse the existing approve route.",
        "Send the dealer a welcome message with login details after approval.",
    ]
    for item in checklist:
        doc.add_paragraph("☐  " + item, style="List Paragraph")

    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run("Note: This is a future plan / idea document. No code has been changed.")
    nr.italic = True
    nr.font.size = Pt(10)
    nr.font.color.rgb = RGBColor.from_string("888888")

    doc.save(OUTPUT_DOCX)


# ---------------------------------------------------------------------------
# 3) Run + self-check
# ---------------------------------------------------------------------------
def main():
    tmp_png = os.path.join(tempfile.gettempdir(), "wa_onboarding_flowchart.png")
    build_flowchart(tmp_png)
    build_doc(tmp_png)

    # Self-check: re-open and validate
    check = Document(OUTPUT_DOCX)
    headings = [p.text for p in check.paragraphs if p.style.name.startswith("Heading")]
    n_images = len(check.inline_shapes)
    size_kb = os.path.getsize(OUTPUT_DOCX) / 1024

    assert n_images >= 1, "Flowchart image missing from the document!"
    assert any("two services" in h for h in headings), "Expected section heading missing!"

    print("OK  ->", OUTPUT_DOCX)
    print(f"    size: {size_kb:.0f} KB")
    print(f"    inline images: {n_images}")
    print(f"    headings: {len(headings)}")
    try:
        os.remove(tmp_png)
    except OSError:
        pass


if __name__ == "__main__":
    main()
